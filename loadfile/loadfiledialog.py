import os.path
import xml.etree.ElementTree as xml

import docx
from PyQt5 import QtWidgets, QtCore, Qt
from PyQt5.QtGui import QFontMetrics
from PyQt5.QtWidgets import QDialog, QMessageBox

from loadfile.loadfilepage_ui import Ui_LoadFileDialog


class LoadFileDialog(QDialog):
    def __init__(self, parent=None):
        super(LoadFileDialog, self).__init__(parent)
        self.ui = Ui_LoadFileDialog()
        self.ui.setupUi(self)
        # self.setAcceptDrops(True)
        self.file_name = None
        self.msg_box = None
        self.success = False
        self.block = False
        self.ui.progressBar.setVisible(False)
        self.PATH_TO_TESTS = "All tests"
        self.PROGRAM_DATA_PATH = os.path.join(os.getenv('APPDATA'), 'Test Generator')

        self.ui.load_cur_razdel.hide()
        self.ui.load_cur_vopros.hide()
        self.ui.load_cur_otvet.hide()
        self.ui.progressBar.hide()

        self.ui.chooseFileBtn.clicked.connect(self.load_file_button_click)
        self.ui.closeDlgBtn.clicked.connect(self.close_dlg)

    # User choosing file
    def load_file_button_click(self):
        if not self.block:
            fname_tuple = QtWidgets.QFileDialog.getOpenFileName(self,
                                                                'Выберите файл',
                                                                os.path.join(os.getenv('USERPROFILE'), 'Desktop'),
                                                                "Документ MS Word (*docx)")
            if fname_tuple[0] != '':
                self.file_name = fname_tuple[0]
                self.load_file()

    # Reading Dropped or Choosed file
    def load_file(self):

        if not self.block:
            # Supports only .docx files
            if self.file_name[-5:len(self.file_name)] == ".docx":

                self.block = True

                doc = docx.Document(self.file_name)

                paragraphs_count = len(doc.paragraphs)

                if paragraphs_count == 0:
                    return

                success = self.__write_questions_to_xml_file(paragraphs_count, doc)

                if success:
                    self.ui.stackedWidget.setCurrentWidget(self.ui.page_2)

            # if file extension is differ than .docx show warning message
            else:
                self.msg_box = QMessageBox()
                self.msg_box.setText("Use files with \".docx\" extension only")
                self.msg_box.setWindowTitle("Warning")
                self.msg_box.setIcon(QMessageBox.Warning)
                self.msg_box.show()

    def __write_questions_to_xml_file(self, paragraphs_count, doc):
        i = 0

        self.ui.load_cur_razdel.show()
        self.ui.load_cur_vopros.show()
        self.ui.load_cur_otvet.show()
        self.ui.progressBar.show()

        path_name, extension = os.path.splitext(self.file_name)
        file_name = path_name.split('/')[-1]

        root = xml.Element("naimenovanie_testa")
        root.text = file_name
        razdel = None
        vopros = None
        otvet = None
        question_id_number = 1

        self.ui.progressBar.setVisible(True)
        self.ui.progressBar.setMaximum(paragraphs_count)

        while i < paragraphs_count:

            QtCore.QCoreApplication.processEvents()

            self.ui.progressBar.setValue(i)

            if doc.paragraphs[i].text == '':
                i += 1
                continue

            italic = bold = underline = None

            # Avoiding empty indent None style
            for run in doc.paragraphs[i].runs:
                if run.text != '':
                    italic = run.italic
                    bold = run.bold
                    underline = run.underline
                    break

            # Add section
            if italic and bold and underline is None:
                razdel = xml.Element("razdel")
                razdel.text = str(doc.paragraphs[i].text)
                root.append(razdel)
                question_id_number = 1

                self.cur_label_text(razdel.text, self.ui.load_cur_razdel, 'razdel')

            # Add next question
            if italic is None and bold and underline is None:
                vopros = xml.Element("vopros", question_id=str(question_id_number))
                vopros.text = str(doc.paragraphs[i].text)
                razdel.append(vopros)
                question_id_number += 1

                self.cur_label_text(vopros.text, self.ui.load_cur_vopros, 'vopros')

            # Add correct answers
            if italic is None and bold and underline:
                otvet = xml.SubElement(vopros, "otvet", status="correct")
                otvet.text = str(doc.paragraphs[i].text)

                self.cur_label_text(otvet.text, self.ui.load_cur_otvet, 'otvet')

            # Add other answers of current question
            if italic is None and bold is None and underline is None:
                otvet = xml.SubElement(vopros, "otvet")
                otvet.text = str(doc.paragraphs[i].text)

                self.cur_label_text(otvet.text, self.ui.load_cur_otvet, 'otvet')

            i += 1

        tree = xml.ElementTree(root)

        os.makedirs(self.PROGRAM_DATA_PATH, exist_ok=True)
        os.makedirs(os.path.join(self.PROGRAM_DATA_PATH, self.PATH_TO_TESTS), exist_ok=True)

        save_folder = os.path.join(self.PROGRAM_DATA_PATH, self.PATH_TO_TESTS)

        with open(os.path.join(save_folder, file_name + '.xml'), 'wb') as file:
            tree.write(file)

        return True

    def cur_label_text(self, text, label: QtWidgets.QLabel, label_type: str):

        if label_type == 'razdel':
            text = 'Раздел: ' + text
            self.ui.load_cur_vopros.setText('')
            self.ui.load_cur_otvet.setText('')
        elif label_type == 'vopros':
            text = 'Вопрос: ' + text
            self.ui.load_cur_otvet.setText('')
        else:
            text = 'Ответ: ' + text

        metrix = QFontMetrics(label.font())
        elided_text = metrix.elidedText(text, Qt.Qt.ElideRight, label.width())
        label.setText(elided_text)

    # Закрыть диалог
    def close_dlg(self):
        self.close()

    # def dragEnterEvent(self, a0: QtGui.QDragEnterEvent) -> None:
    #     if a0.mimeData().hasUrls():
    #         a0.accept()
    #     else:
    #         a0.ignore()
    #
    # def dragMoveEvent(self, a0: QtGui.QDragMoveEvent) -> None:
    #     if a0.mimeData().hasUrls():
    #         a0.accept()
    #     else:
    #         a0.ignore()
    #
    # def dropEvent(self, event: QtGui.QDropEvent) -> None:
    #     if self.file_name is not None:
    #         self.file_name = None
    #
    #     if event.mimeData().hasUrls():
    #         event.setDropAction(QtCore.Qt.CopyAction)
    #         event.accept()
    #         for url in event.mimeData().urls():
    #             self.file_name = str(url.toLocalFile())
    #             break
    #
    #         self.load_file()
    #     else:
    #         event.ignore()
